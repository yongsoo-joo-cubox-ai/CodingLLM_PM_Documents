"""
Secern AI Report Template DOCX Builder

Secern AI 오렌지 브랜딩을 적용한 보고서 DOCX 템플릿을 생성한다.
python-docx를 사용하며, 표지/헤더/푸터/본문 스타일/테이블을 포함한다.

Usage:
    python build_docx_template.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ---------------------------------------------------------------------------
# Secern AI Brand Colors (Orange Branding)
# ---------------------------------------------------------------------------

COLOR_PRIMARY = "ff4800"        # 오렌지 - Heading 1, 테이블 헤더
COLOR_SECONDARY = "b73d0c"      # 진한 오렌지 - Heading 2
COLOR_DARK_TEXT = "333333"       # 본문 텍스트, Heading 3
COLOR_WHITE = "FFFFFF"
COLOR_TABLE_EVEN = "fff5f0"     # 연한 오렌지 - 테이블 짝수행
COLOR_TABLE_ODD = "f8f9fa"      # 연한 회색 - 테이블 홀수행
COLOR_BORDER = "CCCCCC"

# Font stack
FONT_PRIMARY = "Pretendard"
FONT_FALLBACK = "Segoe UI"
FONT_FALLBACK_KO = "Malgun Gothic"
FONT_STACK = [FONT_PRIMARY, FONT_FALLBACK, FONT_FALLBACK_KO]


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """6자리 hex 문자열을 RGBColor로 변환한다."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    """테이블 셀 배경색을 XML로 직접 설정한다."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, color: str = COLOR_BORDER, width: str = "4") -> None:
    """테이블 셀에 얇은 테두리를 설정한다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), width)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _set_paragraph_bottom_border(paragraph, color: str, width: str = "12") -> None:
    """단락 하단에 색상 테두리를 추가한다 (Heading 1 스타일용)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _add_page_number_to_footer(footer) -> None:
    """푸터에 PAGE 필드 코드를 XML로 삽입하여 페이지 번호를 표시한다."""
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Page " 텍스트
    run_prefix = paragraph.add_run("Page ")
    run_prefix.font.size = Pt(9)
    run_prefix.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
    _apply_font_stack(run_prefix)

    # PAGE 필드 삽입
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    run_elem = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run_elem.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = "1"
    run_elem.append(text_elem)
    fld_simple.append(run_elem)
    paragraph._p.append(fld_simple)


def _apply_font_stack(run, font_name: str | None = None) -> None:
    """run에 폰트 스택(Pretendard > Segoe UI > Malgun Gothic)을 설정한다."""
    name = font_name or FONT_PRIMARY
    run.font.name = name
    # East Asian 폰트 설정 (한글 지원)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), FONT_FALLBACK_KO)
    r_fonts.set(qn("w:cs"), FONT_FALLBACK)


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

def _configure_styles(doc: Document) -> None:
    """문서 스타일(Heading 1/2/3, Normal)을 Secern AI 브랜딩으로 설정한다."""
    styles = doc.styles

    # Normal (본문)
    style_normal = styles["Normal"]
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
    style_normal.font.name = FONT_PRIMARY
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15
    # East Asian 폰트
    rpr = style_normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_PRIMARY)
    r_fonts.set(qn("w:hAnsi"), FONT_PRIMARY)
    r_fonts.set(qn("w:eastAsia"), FONT_FALLBACK_KO)
    r_fonts.set(qn("w:cs"), FONT_FALLBACK)

    # Heading 1: 오렌지, 16pt, bold
    style_h1 = styles["Heading 1"]
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = _hex_to_rgb(COLOR_PRIMARY)
    style_h1.font.name = FONT_PRIMARY
    style_h1.paragraph_format.space_before = Pt(18)
    style_h1.paragraph_format.space_after = Pt(8)

    # Heading 2: 진한 오렌지, 14pt, bold
    style_h2 = styles["Heading 2"]
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.font.color.rgb = _hex_to_rgb(COLOR_SECONDARY)
    style_h2.font.name = FONT_PRIMARY
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(6)

    # Heading 3: 어두운 회색, 12pt, bold
    style_h3 = styles["Heading 3"]
    style_h3.font.size = Pt(12)
    style_h3.font.bold = True
    style_h3.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
    style_h3.font.name = FONT_PRIMARY
    style_h3.paragraph_format.space_before = Pt(12)
    style_h3.paragraph_format.space_after = Pt(4)

    # Heading 스타일에도 East Asian 폰트 적용
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        s = styles[style_name]
        rpr = s.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        rf.set(qn("w:ascii"), FONT_PRIMARY)
        rf.set(qn("w:hAnsi"), FONT_PRIMARY)
        rf.set(qn("w:eastAsia"), FONT_FALLBACK_KO)
        rf.set(qn("w:cs"), FONT_FALLBACK)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _configure_page(doc: Document) -> None:
    """A4 용지, 1인치 마진을 설정한다."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------

def _configure_header_footer(doc: Document, doc_number: str) -> None:
    """문서 헤더(기밀 표시)와 푸터(페이지 번호)를 설정한다."""
    section = doc.sections[0]

    # 첫 페이지 헤더/푸터를 다르게 설정 (표지에는 헤더 미표시)
    section.different_first_page_header_footer = True

    # --- 기본(2페이지 이후) 헤더 ---
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(f"{doc_number} | Secern AI Confidential")
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
    header_run.font.italic = True
    _apply_font_stack(header_run)

    # --- 기본(2페이지 이후) 푸터: 페이지 번호 ---
    footer = section.footer
    footer.is_linked_to_previous = False
    _add_page_number_to_footer(footer)

    # --- 첫 페이지 헤더/푸터는 비워둠 (표지) ---
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    if first_header.paragraphs:
        first_header.paragraphs[0].text = ""

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].text = ""


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, doc_number: str) -> None:
    """표지 페이지를 추가한다."""
    # 상단 여백 (표지 중앙 배치용)
    for _ in range(6):
        doc.add_paragraph("")

    # 문서 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("[문서 제목]")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = _hex_to_rgb(COLOR_PRIMARY)
    _apply_font_stack(title_run)

    # 오렌지 구분선 (빈 단락에 하단 테두리)
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_bottom_border(separator, COLOR_PRIMARY, width="18")

    # 빈 줄
    doc.add_paragraph("")

    # 메타데이터 테이블 (표지 중앙)
    meta_items = [
        ("문서번호", doc_number),
        ("작성일", "2026년 X월 X일"),
        ("보안등급", "대외비"),
        ("작성", "Secern AI"),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_idx, (label, value) in enumerate(meta_items):
        # 라벨 셀
        cell_label = table.cell(row_idx, 0)
        cell_label.text = ""
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = _hex_to_rgb(COLOR_SECONDARY)
        _apply_font_stack(run)

        # 값 셀
        cell_value = table.cell(row_idx, 1)
        cell_value.text = ""
        p = cell_value.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
        _apply_font_stack(run)

    # 메타데이터 테이블 테두리 제거
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    tbl_pr.append(borders)

    # 페이지 나누기
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Branded table builder
# ---------------------------------------------------------------------------

def _add_branded_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Secern AI 브랜딩이 적용된 테이블을 추가한다.

    - 헤더행: #ff4800 배경, 흰색 텍스트
    - 짝수행: #fff5f0 연한 오렌지
    - 홀수행: #f8f9fa 연한 회색
    """
    n_cols = len(headers)
    n_rows = len(rows) + 1  # 헤더 포함
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # --- 헤더행 ---
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        _set_cell_shading(cell, COLOR_PRIMARY)
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = _hex_to_rgb(COLOR_WHITE)
        _apply_font_stack(run)

    # --- 데이터행 ---
    for row_idx, row_data in enumerate(rows):
        bg_color = COLOR_TABLE_EVEN if row_idx % 2 == 0 else COLOR_TABLE_ODD
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ""
            _set_cell_shading(cell, bg_color)
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = _hex_to_rgb(COLOR_DARK_TEXT)
            _apply_font_stack(run)

    doc.add_paragraph("")  # 테이블 후 간격


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_template() -> Path:
    """Secern AI Report Template DOCX를 생성하여 저장한다."""
    doc_number = "SAI-XXX-2026-001"
    doc = Document()

    # 1. 스타일 설정
    _configure_styles(doc)

    # 2. 페이지 설정
    _configure_page(doc)

    # 3. 헤더/푸터 설정
    _configure_header_footer(doc, doc_number)

    # 4. 표지 페이지
    _add_cover_page(doc, doc_number)

    # -----------------------------------------------------------------------
    # 5. 본문 샘플 콘텐츠
    # -----------------------------------------------------------------------

    # Heading 1 + 하단 오렌지 테두리
    h1 = doc.add_heading("1. 개요", level=1)
    _set_paragraph_bottom_border(h1, COLOR_PRIMARY)

    p1 = doc.add_paragraph(
        "이 문서는 Secern AI 보고서 템플릿의 표준 양식입니다. "
        "모든 공식 문서는 이 템플릿을 기반으로 작성하며, "
        "일관된 브랜딩과 가독성을 유지합니다."
    )
    for run in p1.runs:
        _apply_font_stack(run)

    # Heading 2
    doc.add_heading("1.1 목적", level=2)

    p2 = doc.add_paragraph(
        "본 템플릿의 목적은 Secern AI의 오렌지 브랜딩 컬러를 "
        "보고서 전체에 일관되게 적용하는 것입니다. "
        "표지, 머리글/바닥글, 제목 스타일, 테이블 서식이 포함됩니다."
    )
    for run in p2.runs:
        _apply_font_stack(run)

    # Heading 3
    doc.add_heading("1.1.1 상세", level=3)

    p3 = doc.add_paragraph(
        "상세 내용은 각 섹션별로 기술하며, "
        "Heading 3 이하의 세부 항목을 통해 구조를 갖춥니다. "
        "본문 텍스트는 Pretendard 11pt, #333333 색상을 사용합니다."
    )
    for run in p3.runs:
        _apply_font_stack(run)

    # -----------------------------------------------------------------------
    # 6. 샘플 테이블
    # -----------------------------------------------------------------------

    h_table = doc.add_heading("2. 데이터 테이블 예시", level=1)
    _set_paragraph_bottom_border(h_table, COLOR_PRIMARY)

    p_table_desc = doc.add_paragraph(
        "아래 테이블은 Secern AI 브랜딩이 적용된 표준 테이블 양식입니다."
    )
    for run in p_table_desc.runs:
        _apply_font_stack(run)

    _add_branded_table(
        doc,
        headers=["구분", "내용", "비고"],
        rows=[
            ["항목 A", "첫 번째 데이터 설명", "완료"],
            ["항목 B", "두 번째 데이터 설명", "진행 중"],
            ["항목 C", "세 번째 데이터 설명", "예정"],
        ],
    )

    # -----------------------------------------------------------------------
    # 7. 변경이력
    # -----------------------------------------------------------------------

    doc.add_paragraph("")  # 간격

    # 구분선
    separator = doc.add_paragraph()
    _set_paragraph_bottom_border(separator, COLOR_BORDER, width="6")

    h_history = doc.add_heading("변경이력", level=1)
    _set_paragraph_bottom_border(h_history, COLOR_PRIMARY)

    _add_branded_table(
        doc,
        headers=["버전", "일자", "변경 내용", "작성자"],
        rows=[
            ["1.0", "2026-XX-XX", "초안 작성", "분석팀"],
        ],
    )

    # -----------------------------------------------------------------------
    # 저장
    # -----------------------------------------------------------------------

    output_path = Path(__file__).resolve().parent / "Secern_AI_Report_Template.docx"
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    output = build_template()
    file_size = output.stat().st_size
    print(f"출력 경로: {output}")
    print(f"파일 크기: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
